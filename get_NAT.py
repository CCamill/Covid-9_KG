'''提取核酸检测结果'''

import docx
import re

def get_NAT(docx_path):
    doc = docx.Document(docx_path)
    start_idx = -1
    paragraphs = []
    for para in doc.paragraphs:
        paragraphs.append(para.text)
        # print(para.text)
    for idx,para in enumerate(paragraphs):
        if '检测情况'  in para and '核酸'  in para and ('三、'  in para or '四、'  in para):
            print(para)
            start_idx = idx+1
            # print(start_idx)
            break

    nat_para = []
    paragrapg = doc.paragraphs[start_idx:]
    if start_idx != -1:
        for para in doc.paragraphs[start_idx:]:
            if '核酸' or '单采' in para.text:
                nat_para.append(para.text)
    print(nat_para)
    positive_text = ''
    nat = '阴性'
    for para in nat_para:
        if ('阳性' in para or '异常' in para) and ('月' in para or '日' in para):
            positive_text = para

        if positive_text != '':
            p = u'[a-z0-9]+月[a-z0-9]{0,2}'
            reg = re.compile(p)
            res = reg.findall(positive_text)
            nat = res[0] + '日 ' + '阳性'

    return nat

if __name__ == '__main__':
    path = r'document/12.04/关于刘善印的调查报告.docx'
    nat = get_NAT(path)
    print(nat)
