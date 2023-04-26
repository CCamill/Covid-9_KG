# 获取个人信息

import docx
import re
from from_id_get_information import get_information_by_id
from get_NAT import get_NAT
from get_address import get_address
from get_phonenumber import get_phonenumber
cn_sym = ['，','。','、','、','；','：','？']

def formalize_filename(filename):
    return re.sub(r'[0-9]+', '', filename)


def get_name(text):
    PATTERN = u'(?:于)[\u4e00-\u9fa5]{2,4}(?:的)'
    reg = re.compile(PATTERN)
    name_list = str(reg.findall(text))
    name = name_list[3:len(name_list) - 3]
    return name


def get_member_name(text):
    PATTERN = u'(?:：)[\u4e00-\u9fa5]{2,11}(?:，男)'
    reg = re.compile(PATTERN)
    name_list = str(reg.findall(text))
    name = name_list[3:len(name_list) - 3]
    return name

def get_health_state(docx_path):
    doc = docx.Document(docx_path)
    text = ''
    for para in doc.paragraphs[2:6]:
        if '健康状况' in para.text:
            text = para.text
    if text != '':
        text = text.replace('：',':')
        text = text.replace(' ','')
        text = text.replace('，',',')
        text = text.replace('、',',')
        print('原文本: ',text)
        index = text.index(':')
        state = text[index + 1:]
        if state != '':
            if state[-1] in cn_sym:
                return state[0:-1]
            else:
                return state
        else:
            return '未知'
    else:
        return '未知'

# def get_NAT(docx_path):
#     doc = docx.Document(docx_path)
#     start_idx = -1
#     for idx,para in enumerate(doc.paragraphs):
#         if '检测情况' and '核酸' and ('三、' or '四、') in para.text:
#             start_idx = idx+1
#
#     nat_para = []
#     if start_idx != -1:
#         for para in doc.paragraphs[start_idx:]:
#             if '核酸' in para.text:
#                 nat_para.append(para.text)
#     positive_text = ''
#     for para in nat_para:
#         if ('阳性' or '异常') and ('月' or '日') in para:
#             positive_text = para
#
#     if positive_text != '':
#         positive_text = positive_text.replace(' ', '')
#         # positive_text = positive_text.replace(',','。').split('。')
#
#         p = u'[a-z0-9]+月[a-z0-9日]{1,2}'
#         reg = re.compile(p)
#         res = reg.findall(positive_text)
#         nat = res[0] + '日 ' + '阳性'
#     else:
#         nat = '阴性'
#     return nat

def get_basic_inf(text):
    inf = {}
    if text == '':
        print('未找到此人基本信息')
        inf['性别'] = '未知'
        inf['年龄'] = '未知'
        inf['身份证号'] = '未知'
    else:
        try:
            reg = re.compile(u'\d{17}[\dXx]')
            id = reg.findall(text)[0]
            inf['身份证号'] = id
        except Exception as e:
            print('未匹配到身份证号或身份证号不准确')
            inf['身份证号'] = '未知'
        demo = get_information_by_id(inf['身份证号'])

        if inf['身份证号'] != '未知' and demo.isTruth():
            inf['年龄'] = demo.getAge()
            inf['性别'] = demo.getSex()
        else:
            reg = re.compile(u'男|女')
            sex = reg.findall(text)[0]
            inf['性别'] = sex

            reg = re.compile(u'\d+岁')
            try:
                age = reg.findall(text)[0]
                inf['年龄'] = age
            except Exception as e:
                try:
                    PATTERN = u'(?:年龄:)[0-9]{2,3}'
                    reg = re.compile(PATTERN)
                    age = reg.findall(text)[0]
                    inf['年龄'] = age
                except Exception as e:
                    inf['年龄'] = '未知'

        # try:
        #     reg = re.compile(u'(?:电话.*)+\d{6,11}')
        #     phone = reg.findall(text)[0]
        #     try:
        #         reg = re.compile('1\d{10}')
        #         phoneNumber = reg.findall(phone)[0]
        #         inf['联系电话'] = phoneNumber
        #     except Exception as e:
        #         reg = re.compile('\d{7}')
        #         phoneNumber = reg.findall(phone)[0]
        #         inf['联系电话'] = phoneNumber
        # except Exception as e:
        #     inf['联系电话'] = '未知'

    return inf


def get_information(docPath):

    doc = docx.Document(docPath)
    paragraphs = []
    for para in doc.paragraphs:
        paragraphs.append(para.text)

    patient_name = get_name(docPath)

    individual = ''

    if patient_name in paragraphs[3]:
        individual = paragraphs[3]  # str
    elif patient_name in paragraphs[4]:
        individual = paragraphs[4]  # str
    elif patient_name in paragraphs[2]:
        individual = paragraphs[2]  # str
    else:
        print('未找到此人基本信息')

    individual = individual.replace('：', ':')
    individual = individual.replace(' ', '')
    individual = individual.replace('，', ',')

    inf = get_basic_inf(individual)

    inf['姓名'] = patient_name
    inf['健康状况'] = get_health_state(docPath)
    inf['核酸检测结果'] = get_NAT(docPath)
    inf['家庭住址'] = get_address(docPath)
    inf['联系电话'] = get_phonenumber(docPath)
    return inf